from odoo import api, fields, models , _
from odoo.tools import float_is_zero, format_amount, format_date, html_keep_url, is_html_empty
from dateutil.relativedelta import relativedelta
from datetime import datetime, timedelta , date
from odoo.exceptions import ValidationError

class SaleOrder(models.Model):
    _inherit = 'sale.order'

    paid_land_1 = fields.Float(compute='get_amounts_paid_land', string='Ene')
    paid_land_2 = fields.Float(compute='get_amounts_paid_land', string='Feb')
    paid_land_3 = fields.Float(compute='get_amounts_paid_land', string='Mar')
    paid_land_4 = fields.Float(compute='get_amounts_paid_land', string='Abr')
    paid_land_5 = fields.Float(compute='get_amounts_paid_land', string='May')
    paid_land_6 = fields.Float(compute='get_amounts_paid_land', string='Jun')
    paid_land_7 = fields.Float(compute='get_amounts_paid_land', string='Jul')
    paid_land_8 = fields.Float(compute='get_amounts_paid_land', string='Agos')
    paid_land_9 = fields.Float(compute='get_amounts_paid_land', string='Sep')
    paid_land_10 = fields.Float(compute='get_amounts_paid_land', string='Oct')
    paid_land_11 = fields.Float(compute='get_amounts_paid_land', string='Nov')
    paid_land_12 = fields.Float(compute='get_amounts_paid_land', string='Dic')
    credit_year_now = fields.Float(compute='get_amounts_paid_land', string='Credito Anual')
    payment_year_now = fields.Float(compute='get_amounts_paid_land', string='Aportado Anual')
    saldo_year_now = fields.Float(compute='get_amounts_paid_land', string='Saldo Anual')


    @api.depends('schedule_land_ids')
    def get_amounts_paid_land(self):
        year = fields.Datetime.now().year
        for record in self:
            record.paid_land_1 = None
            record.paid_land_2 = None
            record.paid_land_3 = None
            record.paid_land_4 = None
            record.paid_land_5 = None
            record.paid_land_6 = None
            record.paid_land_7 = None
            record.paid_land_8 = None
            record.paid_land_9 = None
            record.paid_land_10 = None
            record.paid_land_11 = None
            record.paid_land_12 = None

            credit_year_now = 0
            payment_year_now = 0
            saldo_year_now = 0

            schedule_land_dues = record.get_schedule_x_year(year)

            for sche in schedule_land_dues:
                datex = sche.date

                if datex and datex.year == year:
                    # if datex.month == 1 :
                    #    record.paid_land_1 = sche.amount_due_land

                    if sche.amount_due_land > 0:

                        pagadox = sche.amount_due_land + sche.get_value_adelantos()
                        record[f'paid_land_{datex.month}'] = pagadox
                        payment_year_now += pagadox
                    else:
                        record[f'paid_land_{datex.month}'] = -1 * sche.amount
                        saldo_year_now += sche.amount

                    credit_year_now += sche.amount

            record.credit_year_now = credit_year_now
            record.payment_year_now = payment_year_now
            record.saldo_year_now = saldo_year_now

    def update_credit_saldo(self):
        for record in self:
            total_payment = 0
            total_independence = 0

            for line in record.schedule_land_ids:
                if line.type_schedule in ['dues','advances']:
                    total_payment += line.amount_due_land

                if line.type_schedule in ['independence']:
                    total_independence += line.amount_due_land

            record.total_payment_land = round(total_payment, 2)
            record.saldo_payment_land = round(record.price_credit_land - total_payment, 2)

            record.total_independence_land = round(total_independence, 2)
            record.saldo_independence_land = round(record.price_independence_land - total_independence, 2)

            record.saldo_total_land = record.saldo_independence_land + record.saldo_payment_land

    def get_schedule_x_year(self,year):
        start_date = date(year, 1, 1)
        end_date = date(year, 12, 31)
        # Filtrar los registros
        schedule_land_dues = self.env['schedule.dues.land'].search([
            ('type_schedule', 'in', ['dues']),
            ('order_id', '=', self.id),
            ('date', '>=', start_date),
            ('date', '<=', end_date),
        ])

        return schedule_land_dues

    def reemplazar_parrafo(self, parrafo, reemplazar_dict):
        if any(key in parrafo.text for key in reemplazar_dict.keys()):
            for run in parrafo.runs:
                for buscar, item in reemplazar_dict.items():
                    if buscar in run.text and item['value']:
                        run.text = run.text.replace(buscar, item['value'])

    def reemplazar_texto_plantilla_land(self, doc, reemplazar_dict):

        for shape in doc.inline_shapes:
            if shape.type == 3:
                for parrafo in shape.text_frame.paragraphs:
                    self.reemplazar_parrafo(parrafo,reemplazar_dict)


        # Reemplazar en cuadros de texto
        for shape in doc.inline_shapes:
            for parrafo in shape.text_frame.paragraphs:
                self.reemplazar_parrafo(parrafo,reemplazar_dict)

        # Reemplazar en encabezados
        #for footer in section.footer.paragraphs:
        for section in doc.sections:
            for header in section.header.paragraphs:
                self.reemplazar_parrafo(header,reemplazar_dict)



        for parrafo in doc.paragraphs:

            self.reemplazar_parrafo(parrafo,reemplazar_dict)


    def generar_contrato(self):

        if not self.documents_document_land_id:
            return

        import subprocess
        import sys
        from io import BytesIO

        def install(package):
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

        try:
            from docx import Document
        except:
            install('python-docx')

        try:
            import base64
        except:
            install('base64')

        try:
            from num2words import num2words
        except:
            install('num2words')

        def numero_a_letras(num,formateo=True):
           LABEL_CURRENCY = self.currency_id.currency_unit_label.upper()
           # Convertir el número a letras en español
           if isinstance(num, float):
               entero_part = int(num)
               decimal_part = int(str(num).split(".")[1])
               entero_str = num2words(entero_part, lang='es')
               entero_str = entero_str.upper()
               if not decimal_part:
                   if formateo:
                       return f'''{entero_str} CON 00/100 {LABEL_CURRENCY}'''
                   return entero_str

               if formateo:
                   decimal_part = decimal_part
                   return f'''{entero_str} CON {decimal_part}/100 {LABEL_CURRENCY}'''

               #decimal_str = num2words(decimal_part, lang='es')
               return entero_str
           else:
               entero_str = num2words(num, lang='es')
               entero_str = entero_str.upper()

               if formateo:
                   return f'''{entero_str} CON 00/100 {LABEL_CURRENCY}'''

               return entero_str



        def format_text_currency_contrato(number):
            tformat = "{:,.2f}".format(number)
            return f'''{self.currency_id.symbol} {tformat}'''

        attachment = self.documents_document_land_id.attachment_id
        if not attachment:
            raise ValueError("Attachment no encontrado")

        file_content = base64.b64decode(attachment.datas)

        # Crear un objeto Document a partir del contenido
        doc = Document(BytesIO(file_content))

        marital = self.partner_id.marital

        initial_val = round(self.price_initial_land,2)
        initial_val_format = format_text_currency_contrato(initial_val)

        credit_val = round(self.price_credit_land,2)
        credit_val_format = format_text_currency_contrato(credit_val)

        area_val = round(self.area_lot_related,2)
        area_val_format = format_text_currency_contrato(area_val)

        price_total_val = round(self.price_total_land,2)
        price_total_val_format = format_text_currency_contrato(price_total_val)

        cuota_men_val = round(self.value_due_land,2)
        cuota_men_val_format = format_text_currency_contrato(cuota_men_val)

        num_cuotas = int(self.dues_land)

        if not  self.partner_id.function:
            raise ValidationError('INDIQUE UNA OCUPACION EN EL CLIENTE')

        if not  self.partner_id.marital:
            raise ValidationError('INDIQUE UN ESTADO CIVIL EN EL CLIENTE')

        if not self.partner_id.contact_address_inline:
            raise ValidationError('INDIQUE UNA DIRECCION EN EL CLIENTE')

        if not self.partner_id.l10n_pe_district_name:
            raise ValidationError('INDIQUE UN DISTRITO EN EL CLIENTE')


        values_reemplace = {
            '{{CLIENTE}}' :                 {'value': self.partner_id.display_name } ,
            '{{CLIENTE_DNI}}':              {'value': self.partner_id.vat } ,
            '{{CLIENTE_OCUPACION}}':        {'value': self.partner_id.function} ,
            '{{CLIENTE_ESTADO_CIVIL}}':     {'value': self.partner_id.get_values_marital()[marital] if marital else None } ,
            '{{CLIENTE_DIRECCION}}':        {'value': self.partner_id.contact_address_inline} ,
            '{{CLIENTE_DISTRITO}}':         {'value': self.partner_id.l10n_pe_district_name} ,
            '{{NUMERO_DE_LOTE}}':           {'value': str(self.lot_land)} ,
            '{{LETRA_DE_MANZANA}}':         {'value': self.mz_land} ,
            '{{METRAJE}}':                  {'value': str(area_val_format) } ,
            '{{PRECIO_EN_NUMEROS}}':        {'value': str(price_total_val_format)} ,
            '{{PRECIO_EN_LETRAS}}':         {'value': str(numero_a_letras(price_total_val))} ,
            '{{CUOTA_INICIAL_EN_NUMEROS}}': {'value': str(initial_val_format)} ,
            '{{CUOTA_INICIAL_EN_LETRAS}}':  {'value': str(numero_a_letras(initial_val))} ,
            '{{SALDO_EN_NUMEROS}}':         {'value': str(credit_val_format)} ,
            '{{SALDO_EN_LETRAS}}':          {'value': str(numero_a_letras(credit_val))} ,
            '{{PLAZO_EN_NUMEROS}}':         {'value': str(num_cuotas)} ,
            '{{PLAZO_EN_LETRAS}}':          {'value': str(numero_a_letras(num_cuotas,formateo=False))} ,
            '{{CUOTA_MENSUAL_EN_NUMEROS}}': {'value': str(cuota_men_val_format)} ,
            '{{CUOTA_MENSUAL_EN_LETRAS}}':  {'value': str(numero_a_letras(cuota_men_val))} ,
            '{{EXP}}'                    :  {'value': self.nro_internal_land} ,
            #
        }

        # Reemplazar variables en el documento
        self.reemplazar_texto_plantilla_land(doc, values_reemplace)
        #self.reemplazar_texto_plantilla_land(doc, '{{FECHA}}', '25 de septiembre de 2025')
        #{{CLIENTE_DIRECCION}}

        # Guardar el nuevo documento en un BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        self.contrato_generado_land = base64.b64encode(output.read())
        self.name_contrato_generado_land = f'''CONTRATO_{str(self.nro_internal_land)}_{self.partner_id.name}.docx'''

    def open_product_land(self):
        return {
            "name": f"AGREGAR TERRENO",
            "type": "ir.actions.act_window",
            "view_mode": "form",
            "view_id": self.env.ref('land.add_terreno_sale').id,
            "res_model": "sale.order",
            "res_id": self.id,
            "target": "new",
            #"context": {
            #    'default_order_id': self.id
            #}

        }

    def update_dates_land(self):
        for invc in self.invoice_ids:
            if invc.amount_total == self.price_initial_land:
                invc.invoice_date = self.date_sign_land

            if invc.state == 'draft':
                invc.action_post()
            #    exist_confirm = True

        if self.price_total_land and self.price_total_land != 0 and len(
                self.invoice_ids) > 1 and self.date_first_due_land:
            invoices = self.env['account.move'].search([
                ('id', 'in', self.invoice_ids.ids),
            ], order='invoice_date asc')
            date_init = self.date_first_due_land
            is_end_month = False
            if date_init.day > 25 and date_init.day <= 31:
                is_end_month = True

            c = 0
            for invoice in invoices:

                if c > 0:
                    if is_end_month:
                        cx = c - 1
                        if cx > 0 :
                            date_initx = date_init + relativedelta(months=cx)
                        else:
                            date_initx = date_init




                        last_date = datetime(date_initx.year if date_initx.month != 12 else date_initx.year + 1,
                                             date_initx.month + 1 if date_initx.month != 12 else 1, 1) - timedelta(
                            days=1)

                        if last_date.day != date_init.day:
                            date_initx = last_date

                    else:
                        cx = c - 1
                        if cx > 0:
                            date_initx = date_init + relativedelta(months=cx)
                        else:
                            date_initx = date_init

                    invoice.invoice_date = date_initx
                c += 1

    def show_dues_land(self):
        self.update_schedule()
        return {
            "name": f"PAGOS",
            "type": "ir.actions.act_window",
            "view_mode": "form",
            "view_id": self.env.ref('land.view_order_form_due').id ,
            "res_model": "sale.order",
            "res_id": self.id,
            "target": "new",

        }

    def show_lot_availables(self):
        product = self.env['product.template'].search([('payment_land_dues','=',True),('sale_ok','=',True)])
        product.update_lots_jz()

        return {
            "name": f"LOTES",
            "type": "ir.actions.act_window",
            "view_mode": "kanban,tree",
            # "view_id": self.env.ref('land.view_order_form_due').id,
            "res_model": "report.lot.land.line",
            "res_id": product.id,
            "target": "current",
            "domain": [('product_tmp_id', '=', product.id)],
            "context": {
                'search_default_gr_mz_value_id': 1
            }

        }

    def update_all_info_land(self):
        record = self.env['sale.order'].search([])
        record.get_info_land()

    def recreate_schedule(self):
        for record in self:
            record.schedule_land_ids.unlink()
        self.update_schedule()

    def update_schedule_all(self):
        orders = self.env['sale.order'].search([])
        orders.update_schedule()
        orders.update_credit_saldo()

    def invoice_lines_available_land(self):
        invoice_lines_dues = []
        invoice_lines_initial = []
        invoice_lines_indepen = []


        amount_indepenced = 0
        qty_to_indepenced = 0

        for line in self.order_line:

            #para independizacion
            if line.product_id.is_independence:
                qty_to_indepenced += line.product_uom_qty
                amount_indepenced += line.price_unit

            for line_inv in line.invoice_lines:

                if line_inv.move_id.move_type in ['out_refund']:
                    continue


                if line_inv.move_id.payment_state == 'reversed' or line_inv.move_id.l10n_pe_edi_reversal_type_id:
                    continue

                if line_inv.move_id.debit_origin_id or line_inv.move_id.state == 'cancel':
                    continue


                #para cuotas
                if line.product_id.payment_land_dues and not line.product_id.is_independence:

                    x = range(int(line_inv.quantity))

                    for n in x:
                        invoice_lines_dues.append(line_inv)

                #para iniciales
                if line.product_id.is_advanced_land or line.product_id.is_separation_land:
                    invoice_lines_initial.append(line_inv)




                #para independicacion
                if line.product_id.is_independence:
                    invoice_lines_indepen.append(line_inv)




        if invoice_lines_dues:
            invoice_lines_dues.reverse()

        if invoice_lines_indepen:
            invoice_lines_indepen.reverse()

        return invoice_lines_dues , invoice_lines_initial , qty_to_indepenced , invoice_lines_indepen , amount_indepenced

    def _recalcule_price_land(self):
        for record in self:

            for line in record.order_line:
                if line.product_id and line.product_id.payment_land_dues:

                    line.change_product_uom_qty_land()

    def write(self,values):
        res = super().write(values)

        if  'report_lot_land_line_id' in values and not self.order_line:
            if self.report_lot_land_line_id:

                amount_total = self.report_lot_land_line_id.price

                if self.inicial_lot_set > 0:
                    amount_total -= self.inicial_lot_set

                self.order_line += self.env['sale.order.line'].new({
                    'product_id': self.report_lot_land_line_id.product_tmp_id.product_variant_ids.id,
                    'price_unit': amount_total / self.report_lot_land_line_id.product_tmp_id.dues_qty,
                    'product_uom_qty': self.report_lot_land_line_id.product_tmp_id.dues_qty

                })

            if self.inicial_lot_set and self.report_lot_land_line_id:
                if self.inicial_lot_set > 0:
                    self.order_line += self.env['sale.order.line'].new({
                        'product_id': self.report_lot_land_line_id.product_tmp_id.optional_product_ids[
                            0].product_variant_ids.id,
                        'price_unit': self.inicial_lot_set,
                        'product_uom_qty': 1
                    })




        return res

    # esto verifica si existe una factura de separacion
    def check_adelanto(self):
        for record in self:

            line_set = None
            amount_set = None

            if record.move_separation_land_id:
                if len(record.order_line) == 2:
                    for line in record.order_line:
                        if line.product_id.is_advanced_land:
                            product_id = record.move_separation_land_id.invoice_line_ids[0].product_id

                            clone_line = line.copy(default={
                                'order_id': record.id,
                                'product_id': product_id.id,
                                'tax_id': [(6, 0, product_id.taxes_id.ids)]
                            })
                            clone_line.price_unit = record.move_separation_land_id.amount_total

                            price_unit_new = line.price_unit - record.move_separation_land_id.amount_total
                            line.price_unit = price_unit_new * 1

                            line_set = line
                            amount_set = price_unit_new * 1

                            record.move_separation_land_id.stage_separation_land = 'initial'

                            # raise ValueError(line.price_unit)

            if line_set:
                line_set.price_unit = amount_set

                # raise ValueError(line_set.price_unit)
            # raise ValueError(line_set)

    def _get_invoiceable_lines(self, final=False):

        if self.sale_line_payment_id:
            return self.sale_line_payment_id



        """Return the invoiceable lines for order `self`."""
        down_payment_line_ids = []
        invoiceable_line_ids = []
        pending_section = None
        precision = self.env['decimal.precision'].precision_get('Product Unit of Measure')

        quantity_lines_invoice = 0

        have_separation = False

        for line in self.order_line:
            if line.display_type == 'line_section':
                continue
            if line.display_type != 'line_note' and float_is_zero(line.qty_to_invoice, precision_digits=precision):
                continue

            if line.qty_to_invoice > 0 or (line.qty_to_invoice < 0 and final) or line.display_type == 'line_note':
                if line.is_downpayment:
                    continue

            if line.product_id.is_separation_land:
                have_separation = True

            quantity_lines_invoice += 1


        for line in self.order_line:



            if line.display_type == 'line_section':
                # Only invoice the section if one of its lines is invoiceable
                pending_section = line
                continue
            if line.display_type != 'line_note' and float_is_zero(line.qty_to_invoice, precision_digits=precision):
                continue

            #if quantity_lines_invoice > 1 and line.product_id.payment_land_dues:
            #    continue

            if have_separation and not line.product_id.is_separation_land:
                continue

            if line.qty_to_invoice > 0 or (line.qty_to_invoice < 0 and final) or line.display_type == 'line_note':
                if line.is_downpayment:
                    # Keep down payment lines separately, to put them together
                    # at the end of the invoice, in a specific dedicated section.
                    down_payment_line_ids.append(line.id)
                    continue
                if pending_section:
                    invoiceable_line_ids.append(pending_section.id)
                    pending_section = None
                invoiceable_line_ids.append(line.id)

        res = self.env['sale.order.line'].browse(invoiceable_line_ids + down_payment_line_ids)



        return res

    def _prepare_invoice(self):
        self.get_last_payment_date_land()
        res = super()._prepare_invoice()
        if self.journal_id:
            res['journal_id'] = self.journal_id.id

        if self.days_expired_land:
            res['days_expired_land'] = self.days_expired_land
            res['value_mora_land'] = self.value_mora_land

        res['invoice_payment_term_id'] =  self.env.ref('account.account_payment_term_immediate').id

        return res

    def print_report_schedule_excell(self):
        return self.env['report.schedule.land'].print_report_schedule_excell(self)

    def action_confirm(self):

        self.check_adelanto()
        res = super().action_confirm()
        if self.move_separation_land_id:
            for line in self.order_line:
                for linex in self.move_separation_land_id.invoice_line_ids:
                    if linex.product_id == line.product_id and line.price_unit == linex.price_unit :
                        line.invoice_lines = [(4, linex.id)]
                        #self.move_separation_land_id.is_separation_land = False


        return res