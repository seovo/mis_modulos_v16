from odoo import api, fields, models , _
from odoo.tools import float_is_zero, format_amount, format_date, html_keep_url, is_html_empty
from dateutil.relativedelta import relativedelta
from datetime import datetime, timedelta , date
from odoo.exceptions import ValidationError

class SaleOrder(models.Model):
    _inherit = 'sale.order'

    documents_document_land_id = fields.Many2one('documents.document',string="Contrato Plantilla")
    contrato_generado_land = fields.Binary(string='Contrato Generado')
    name_contrato_generado_land = fields.Char()

    def reemplazar_texto_plantilla_land(self, doc, buscar, reemplazar, negrita=False):
        for parrafo in doc.paragraphs:
            if buscar in parrafo.text:
                # Limpiar el párrafo
                parrafo.clear()

                # Dividir el texto en partes
                partes = parrafo.text.split(buscar)

                # Agregar texto anterior
                if partes[0]:
                    parrafo.add_run(partes[0])

                # Agregar el texto reemplazado con formato
                run = parrafo.add_run(reemplazar)
                if negrita:
                    run.bold = True

                # Agregar texto posterior si existe
                if len(partes) > 1 and partes[1]:
                    parrafo.add_run(partes[1])




    def reemplazar_texto_plantilla_land_old(self, doc, buscar, reemplazar):
        for parrafo in doc.paragraphs:
            if buscar in parrafo.text:
                parrafo.text = parrafo.text.replace(buscar, reemplazar)

    def generar_contrato(self):

        if not self.documents_document_land_id:
            return

        import subprocess
        import sys
        from io import BytesIO

        def install(package):
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

        try:
            from docx import Document
        except:
            install('python-docx')

        try:
            import base64
        except:
            install('base64')

        attachment = self.documents_document_land_id.attachment_id
        if not attachment:
            raise ValueError("Attachment no encontrado")

        file_content = base64.b64decode(attachment.datas)

        # Crear un objeto Document a partir del contenido
        doc = Document(BytesIO(file_content))


        # Reemplazar variables en el documento
        self.reemplazar_texto_plantilla_land(doc, '{{CLIENTE}}', self.partner_id.display_name , negrita=True)
        self.reemplazar_texto_plantilla_land(doc, '{{CLIENTE_DNI}}', self.partner_id.vat)
        self.reemplazar_texto_plantilla_land(doc, '{{CLIENTE_OCUPACION}}', self.partner_id.function)
        self.reemplazar_texto_plantilla_land(doc, '{{CLIENTE_ESTADO_CIVIL}}', self.partner_id.function)
        self.reemplazar_texto_plantilla_land(doc, '{{CLIENTE_DIRECCION}}', self.partner_id.contact_address_inline)

        #self.reemplazar_texto_plantilla_land(doc, '{{DIRECCION}}', 'Av. Siempre Viva 123')
        #self.reemplazar_texto_plantilla_land(doc, '{{FECHA}}', '25 de septiembre de 2025')
        #{{CLIENTE_DIRECCION}}

        # Guardar el nuevo documento en un BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        self.contrato_generado_land = base64.b64encode(output.read())
        self.name_contrato_generado_land = 'contrato_generado.docx'


